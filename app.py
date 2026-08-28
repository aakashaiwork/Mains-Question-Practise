import datetime
import io
import re
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from groq import Groq
from openai import OpenAI
import pandas as pd
import streamlit as st
from streamlit_gsheets import GSheetsConnection

# --- STREAMLIT UI SETUP ---
st.set_page_config(
    page_title="UPSC/GPSC Mains Paper Generator",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# --- HIDE STREAMLIT BRANDING, HEADER & FOOTER ---
hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    header {visibility: hidden;}
    [data-testid="stHeader"] {display: none !important;}
    footer {visibility: hidden;}
    [data-testid="stFooter"] {display: none !important;}
    .stAppDeployButton {display: none !important;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# System Prompt containing multi-language guidelines
SYSTEM_PROMPT = """
You are an exclusive Civil Services Examination Content Creator and Senior Evaluator specializing solely in UPSC (Union Public Service Commission) and GPSC (Gujarat Public Service Commission) Mains General Studies papers.
Your primary objective is to generate high-quality Mains examination questions and model solutions strictly matching the current difficulty, syllabus, analytical depth, and formatting standards of UPSC/GPSC.

# STRICT SCOPE ENFORCEMENT & OUT-OF-SCOPE DECLINATIONS
If the user's input is off-topic, general knowledge trivia, coding, personal advice, or non-exam related:
YOU MUST DECLINE TO ANSWER with: "This application is strictly configured to generate UPSC/GPSC Mains Answer Writing Papers. Please enter a valid civil services subject/topic."

# DIFFICULTY LEVEL & STRICT WORD COUNT FRAMEWORK
Tailor the question length, complexity, and answer word count strictly according to the requested difficulty level:

1. EASY LEVEL:
   - Question Style: Short, direct analytical question focusing on core concepts, foundational policy issues, or simple constitutional provisions.
   - Word Count Target: ~200 words total per answer (Intro: 25-30w | Body: 140-150w | Conclusion: 25-30w).
   - Marks: 10 Marks.

2. MODERATE LEVEL:
   - Question Style: Standard Mains multi-dimensional analytical question combining static theory with current affairs or policy bottlenecks.
   - Word Count Target: ~300 words total per answer (Intro: 35-40w | Body: 220-230w | Conclusion: 35-40w).
   - Marks: 15 Marks.

3. DIFFICULT LEVEL:
   - Question Style: Long, highly nuanced, quote-based, statement-driven, or contemporary policy dilemma question requiring deep multi-disciplinary synthesis.
   - Word Count Target: ~500 words total per answer (Intro: 50-60w | Body: 380-400w | Conclusion: 50-60w).
   - Marks: 20 Marks.

# DOCUMENT FORMATTING SPECIFICATIONS
- Do NOT include any visual placeholders, diagram hints, or textual blocks for flowcharts/diagrams.

# LANGUAGE DIRECTIVE
You will be provided a list of selected target languages. You MUST generate the complete exam paper and model solution sequentially for EACH requested language, creating a separate, clearly labeled section for every language:
- For English: "# 📝 SECTION: ENGLISH VERSION"
- For Gujarati: "# 📝 SECTION: GUJARATI VERSION (ગુજરાતી આવૃત્તિ)"
- For Hindi: "# 📝 SECTION: HINDI VERSION (हिंदी संस्करण)"

Maintain high academic rigor, formal administrative terminology, and native fluency appropriate for civil services examinations in all generated translations.

# MODEL SOLUTION STRUCTURE
1. Introduction (10-15% of Word Limit): Open directly with a definition, recent context/news, Constitutional Article/Supreme Court judgment, relevant statistic, or Committee recommendation.
2. Body (75-80% of Word Limit): Clear sub-headings with concise bullet points and bold lead-ins. Seamlessly include value additions (Data, Articles, Committee Reports, NITI Aayog papers, SDGs).
3. Conclusion (10-15% of Word Limit): Forward-looking, solution-oriented, and constructive (e.g., Viksit Bharat @2047, Net Zero, Constitutional ideals).

# OUTPUT FORMAT FOR EACH LANGUAGE SECTION
## DAILY MAINS ANSWER WRITING PAPER
Target Exam: [UPSC / GPSC] | Subject: [Subject Name] | Difficulty: [Easy / Moderate / Difficult]  
Total Questions: [Count] | Word Limit: [200 / 300 / 500 words per answer]  

### QUESTION 1
[Question text in requested language]  
Marks: [10 / 15 / 20 Marks] | Word Limit: [200 / 300 / 500 words]

#### MODEL ANSWER
1. Introduction  
[Introduction text in requested language]

2. Body  
[Sub-headings and bullet points in requested language]

3. Conclusion  
[Conclusion text in requested language]
"""


# Helper function to convert markdown text to formatted Word (.docx) file
def create_docx(text_content):
  doc = docx.Document()

  # Set Narrow Margins (0.5 inches on all sides)
  for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

  # Set default font style to Verdana Size 10
  style = doc.styles["Normal"]
  font = style.font
  font.name = "Verdana"
  font.size = Pt(10)

  # Process markdown text line by line
  for line in text_content.split("\n"):
    line = line.strip()
    if not line:
      doc.add_paragraph()
      continue

    if line.startswith("# "):
      p = doc.add_paragraph()
      run = p.add_run(line.replace("# ", ""))
      run.font.size = Pt(14)
      run.bold = True
    elif line.startswith("## "):
      p = doc.add_paragraph()
      run = p.add_run(line.replace("## ", ""))
      run.font.size = Pt(12)
      run.bold = True
    elif line.startswith("### "):
      p = doc.add_paragraph()
      run = p.add_run(line.replace("### ", ""))
      run.font.size = Pt(11)
      run.bold = True
    elif line.startswith("- "):
      p = doc.add_paragraph(style="List Bullet")
      p.add_run(line.replace("- ", ""))
    else:
      p = doc.add_paragraph(line)

  buffer = io.BytesIO()
  doc.save(buffer)
  buffer.seek(0)
  return buffer


# --- DATABASE HELPERS (GOOGLE SHEETS) ---
def get_past_questions_from_db(subject_name):
  try:
    if "connections" in st.secrets and "gsheets" in st.secrets["connections"]:
      conn = st.connection("gsheets", type=GSheetsConnection)
      df = conn.read(ttl="0s")
      if df is not None and not df.empty and "Subject" in df.columns:
        matching_rows = df[
            df["Subject"].str.contains(subject_name, case=False, na=False)
        ]
        if not matching_rows.empty and "Question_Text" in matching_rows.columns:
          return matching_rows["Question_Text"].dropna().tolist()
  except Exception:
    pass
  return []


def save_question_to_db(
    target_exam, subject, topic, difficulty, generated_text
):
  try:
    if "connections" in st.secrets and "gsheets" in st.secrets["connections"]:
      conn = st.connection("gsheets", type=GSheetsConnection)
      existing_df = conn.read(ttl="0s")

      question_snippet = generated_text
      if "### QUESTION 1" in generated_text:
        question_snippet = (
            generated_text.split("### QUESTION 1")[1]
            .split("#### MODEL ANSWER")[0]
            .strip()
        )

      new_row = pd.DataFrame([{
          "Date": str(datetime.date.today()),
          "Target_Exam": target_exam,
          "Subject": subject,
          "Topic": topic if topic else "Auto-selected",
          "Difficulty": difficulty,
          "Question_Text": question_snippet[:500],
      }])

      if existing_df is not None and not existing_df.empty:
        updated_df = pd.concat([existing_df, new_row], ignore_index=True)
      else:
        updated_df = new_row

      conn.update(data=updated_df)
  except Exception:
    pass


# --- DUAL-ENGINE GENERATION WITH AUTOMATIC FALLBACK ---
def generate_mains_paper(user_prompt, num_langs, nvidia_key, groq_key):
  max_tokens_val = min(4000 * num_langs, 16384)

  # Priority 1: NVIDIA Nemotron-3 Ultra
  if nvidia_key:
    try:
      client_nv = OpenAI(
          base_url="https://integrate.api.nvidia.com/v1", api_key=nvidia_key
      )
      response = client_nv.chat.completions.create(
          model="nvidia/nemotron-3-ultra-550b-a55b",
          messages=[
              {"role": "system", "content": SYSTEM_PROMPT},
              {"role": "user", "content": user_prompt},
          ],
          temperature=0.6,
          top_p=0.95,
          max_tokens=max_tokens_val,
          extra_body={"chat_template_kwargs": {"enable_thinking": False}},
      )
      raw_output = response.choices[0].message.content
      cleaned_output = re.sub(
          r"<think>.*?</think>", "", raw_output, flags=re.DOTALL
      ).strip()
      return cleaned_output, "NVIDIA Nemotron-3 Ultra"
    except Exception as nv_err:
      st.toast(
          f"Nemotron API unavailable, switching to Groq fallback... ({str(nv_err)[:50]})"
      )

  # Priority 2: Groq Fallback
  if groq_key:
    client_groq = Groq(api_key=groq_key)
    response = client_groq.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.4,
        max_completion_tokens=min(4000 * num_langs, 8000),
    )
    return response.choices[0].message.content, "Groq (GPT OSS 120B)"

  raise ValueError("Neither NVIDIA API Key nor Groq API Key is available.")


# --- MAIN INTERFACE HEADER ---
st.title("📝 UPSC / GPSC Daily Mains Paper Generator")
st.caption(
    "Dual AI Engine (NVIDIA Nemotron-3 Ultra with Groq Fallback) | Integrated"
    " Question Bank"
)

# Fetch API Keys from Streamlit Secrets
nvidia_api_key = st.secrets.get("NVIDIA_API_KEY", "")
groq_api_key = st.secrets.get("GROQ_API_KEY", "")

# Fallback inputs in sidebar if missing from Secrets
if not nvidia_api_key and not groq_api_key:
  st.sidebar.header("⚙️ API Configuration")
  nvidia_api_key = st.sidebar.text_input(
      "NVIDIA API Key (Priority 1)", type="password"
  )
  groq_api_key = st.sidebar.text_input("Groq API Key (Fallback)", type="password")

# --- MAIN FORM INPUTS ---
st.subheader("📋 Paper Parameters")

col_sub, col_top = st.columns(2)

with col_sub:
  subject_input = st.text_input(
      "1. Subject (Required)",
      placeholder="e.g., GS-2 Polity, GS-3 Economy, GS-1 History, Ethics",
  )

with col_top:
  topic_input = st.text_input(
      "2. Topic / Sub-topic (Optional)",
      placeholder=(
          "e.g., Judicial Activism, Inflation (Leave blank for auto-selection)"
      ),
  )

col1, col2, col3, col4 = st.columns([1.5, 1.2, 2.5, 1.2])

with col1:
  difficulty = st.selectbox(
      "3. Difficulty Level",
      ["Moderate", "Easy", "Difficult"],
      help="Easy = ~200 words | Moderate = ~300 words | Difficult = ~500 words",
  )

with col2:
  target_exam = st.selectbox("4. Target Exam", ["UPSC", "GPSC"])

with col3:
  selected_languages = st.multiselect(
      "5. Select Language(s)",
      ["English", "Gujarati", "Hindi"],
      default=["English"],
      help="Check one or multiple languages to generate outputs for.",
  )

with col4:
  num_questions = st.number_input(
      "6. Questions", min_value=1, max_value=3, value=1
  )

st.divider()

# --- GENERATE ACTION ---
if st.button("🚀 Generate Mains Paper", type="primary", use_container_width=True):
  if not nvidia_api_key and not groq_api_key:
    st.error(
        "API Key is missing. Please add NVIDIA_API_KEY or GROQ_API_KEY in"
        " Secrets."
    )
  elif not subject_input.strip():
    st.warning("Please enter a Subject (e.g., GS-2 Polity, GS-3 Economy).")
  elif not selected_languages:
    st.warning("Please select at least one language.")
  else:
    try:
      # 1. Fetch past questions from Google Sheet Database
      past_questions = get_past_questions_from_db(subject_input.strip())
      anti_duplication_prompt = ""

      if past_questions:
        past_q_str = "\n- ".join(past_questions[-10:])
        anti_duplication_prompt = (
            "\n\nCRITICAL ANTI-DUPLICATION INSTRUCTION:\nDo NOT repeat,"
            " rephrase, or generate questions similar to these previously"
            " generated questions from the database:\n- "
            + past_q_str
        )

      # 2. Build Prompt
      if topic_input.strip():
        topic_details = (
            f"Subject: '{subject_input.strip()}', Specific Topic:"
            f" '{topic_input.strip()}'"
        )
        file_name_tag = (
            f"{subject_input}_{topic_input}".replace(" ", "_")
            .replace("/", "_")
            .strip()
        )
      else:
        topic_details = (
            f"Subject: '{subject_input.strip()}'. (Please automatically select a"
            " high-yield, priority topic suitable for Mains from this subject)"
        )
        file_name_tag = (
            subject_input.replace(" ", "_").replace("/", "_").strip()
        )

      languages_str = ", ".join(selected_languages)
      user_prompt = f"Generate a {target_exam} Daily Mains Answer Writing Paper strictly in the following selected language(s): [{languages_str}]. {topic_details}. Difficulty Level: {difficulty}. Total Questions: {num_questions}.{anti_duplication_prompt}"

      with st.spinner(f"Generating paper in ({languages_str})..."):
        generated_paper, engine_used = generate_mains_paper(
            user_prompt=user_prompt,
            num_langs=len(selected_languages),
            nvidia_key=nvidia_api_key,
            groq_key=groq_api_key,
        )

        # Save to DB if connection exists
        save_question_to_db(
            target_exam,
            subject_input.strip(),
            topic_input.strip(),
            difficulty,
            generated_paper,
        )

        st.success(
            f"Mains Paper Generated Successfully in ({languages_str}) via"
            f" {engine_used}!"
        )

        # Download Button
        docx_file = create_docx(generated_paper)
        st.download_button(
            label=(
                "📥 Download as Pre-formatted Word Document (.docx) -"
                f" {languages_str}"
            ),
            data=docx_file,
            file_name=(
                f"{target_exam}_{file_name_tag}_{difficulty}_{'_'.join(selected_languages)}_Paper.docx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            type="primary",
        )

        st.divider()
        st.markdown(generated_paper)

    except Exception as e:
      st.error(f"An error occurred: {str(e)}")
